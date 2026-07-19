import os
import subprocess
import tempfile
from pathlib import Path
import torch
import whisper
from pyannote.audio import Pipeline
from docx import Document
from docx.shared import Inches
import librosa
import numpy as np
from datetime import datetime, timedelta
import threading
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import sys
from dotenv import load_dotenv

load_dotenv()

class VideoTranscriber:
    def __init__(self, progress_callback=None, log_callback=None):
        """Инициализация транскрибера"""
        self.progress_callback = progress_callback
        self.log_callback = log_callback
        
        self.log("Загрузка моделей...")
        
        # Загружаем модель Whisper для транскрибации
        self.whisper_model = whisper.load_model("base")
        
        # Загружаем модель pyannote для разделения спикеров
        hf_token = os.getenv("HUGGINGFACE_TOKEN")
        try:
            if not hf_token:
                raise ValueError("HUGGINGFACE_TOKEN не задан в .env")
            self.diarization_pipeline = Pipeline.from_pretrained(
                "pyannote/speaker-diarization-3.1",
                token=hf_token
            )
        except Exception as e:
            self.log(f"Ошибка загрузки модели диаризации: {e}")
            self.log("Для работы с разделением спикеров нужен токен Hugging Face")
            self.diarization_pipeline = None
            
        self.log("Модели загружены!")

    def log(self, message):
        """Отправка сообщения в лог"""
        if self.log_callback:
            self.log_callback(message)
        else:
            print(message)

    def update_progress(self, value):
        """Обновление прогресс-бара"""
        if self.progress_callback:
            self.progress_callback(value)

    def extract_audio_from_video(self, video_path, output_path):
        """Извлечение аудио из видео файла"""
        try:
            ffmpeg_path = 'ffmpeg'  # Если FFmpeg в PATH
            
            cmd = [
                ffmpeg_path, '-i', video_path, 
                '-ar', '16000', '-ac', '1', 
                '-c:a', 'pcm_s16le', 
                '-y', output_path
            ]
            subprocess.run(cmd, check=True, capture_output=True)
            return True
        except subprocess.CalledProcessError as e:
            self.log(f"Ошибка извлечения аудио: {e}")
            return False
        except FileNotFoundError:
            self.log("FFmpeg не найден. Установите FFmpeg или укажите полный путь в коде.")
            return False

    def perform_speaker_diarization(self, audio_path):
        """Разделение спикеров в аудио файле"""
        if self.diarization_pipeline is None:
            self.log("Модель диаризации недоступна. Будет использован один спикер.")
            return None
            
        try:
            diarization = self.diarization_pipeline(audio_path)
            
            speaker_segments = []
            for turn, _, speaker in diarization.itertracks(yield_label=True):
                speaker_segments.append({
                    'start': turn.start,
                    'end': turn.end,
                    'speaker': speaker
                })
            return speaker_segments
        except Exception as e:
            self.log(f"Ошибка диаризации: {e}")
            return None

    def transcribe_audio(self, audio_path, language='ru'):
        """Транскрибация аудио файла"""
        try:
            result = self.whisper_model.transcribe(
                audio_path,
                language=language,
                word_timestamps=True
            )
            return result
        except Exception as e:
            self.log(f"Ошибка транскрибации: {e}")
            return None

    def merge_transcription_with_speakers(self, transcription, speaker_segments):
        """Объединение транскрибации с информацией о спикерах"""
        if not speaker_segments:
            merged_result = []
            for segment in transcription['segments']:
                merged_result.append({
                    'start': segment['start'],
                    'end': segment['end'],
                    'text': segment['text'].strip(),
                    'speaker': 'Человек 1'
                })
            return merged_result

        merged_result = []
        speaker_mapping = {}
        speaker_counter = 1
        
        for segment in transcription['segments']:
            segment_start = segment['start']
            segment_end = segment['end']
            segment_text = segment['text'].strip()
            
            best_speaker = None
            max_overlap = 0
            
            for speaker_seg in speaker_segments:
                overlap_start = max(segment_start, speaker_seg['start'])
                overlap_end = min(segment_end, speaker_seg['end'])
                overlap = max(0, overlap_end - overlap_start)
                
                if overlap > max_overlap:
                    max_overlap = overlap
                    best_speaker = speaker_seg['speaker']
            
            if best_speaker and best_speaker not in speaker_mapping:
                speaker_mapping[best_speaker] = f"Человек {speaker_counter}"
                speaker_counter += 1
            
            speaker_name = speaker_mapping.get(best_speaker, "Человек 1")
            
            merged_result.append({
                'start': segment_start,
                'end': segment_end,
                'text': segment_text,
                'speaker': speaker_name
            })
        
        return merged_result

    def format_time(self, seconds):
        """Форматирование времени в читаемый вид"""
        td = timedelta(seconds=seconds)
        hours, remainder = divmod(td.seconds, 3600)
        minutes, seconds = divmod(remainder, 60)
        return f"{hours:02d}:{minutes:02d}:{seconds:02d}"

    def create_word_document(self, merged_transcription, output_path, video_name):
        """Создание Word документа с транскрибацией"""
        doc = Document()
        
        title = doc.add_heading(f'Транскрибация: {video_name}', 0)
        
        info_para = doc.add_paragraph()
        info_para.add_run(f"Дата создания: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}\n")
        info_para.add_run(f"Исходный файл: {video_name}\n")
        info_para.add_run(f"Количество спикеров: {len(set(item['speaker'] for item in merged_transcription))}")
        
        doc.add_paragraph("=" * 50)
        
        current_speaker = None
        current_text = ""
        current_start_time = None
        
        for item in merged_transcription:
            if item['speaker'] != current_speaker:
                if current_speaker is not None and current_text.strip():
                    para = doc.add_paragraph()
                    para.add_run(f"{current_speaker}: ").bold = True
                    para.add_run(current_text.strip())
                    
                    time_para = doc.add_paragraph()
                    time_para.add_run(f"[{self.format_time(current_start_time)} - {self.format_time(item['start'])}]").italic = True
                    time_para.alignment = 2
                
                current_speaker = item['speaker']
                current_text = item['text']
                current_start_time = item['start']
            else:
                current_text += " " + item['text']
        
        if current_speaker is not None and current_text.strip():
            para = doc.add_paragraph()
            para.add_run(f"{current_speaker}: ").bold = True
            para.add_run(current_text.strip())
            
            time_para = doc.add_paragraph()
            time_para.add_run(f"[{self.format_time(current_start_time)} - {self.format_time(merged_transcription[-1]['end'])}]").italic = True
            time_para.alignment = 2
        
        doc.save(output_path)
        self.log(f"Документ сохранен: {output_path}")

    def process_video(self, video_path, output_dir=None, language='ru'):
        """Основная функция обработки видео"""
        video_path = Path(video_path)
        
        if not video_path.exists():
            self.log(f"Файл не найден: {video_path}")
            return False
        
        if output_dir is None:
            output_dir = video_path.parent
        else:
            output_dir = Path(output_dir)
            output_dir.mkdir(parents=True, exist_ok=True)
        
        self.log(f"Обработка видео: {video_path.name}")
        
        with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as temp_audio:
            temp_audio_path = temp_audio.name
        
        try:
            # Шаг 1: Извлечение аудио
            self.log("1. Извлечение аудио из видео...")
            self.update_progress(10)
            if not self.extract_audio_from_video(str(video_path), temp_audio_path):
                return False
            
            # Шаг 2: Разделение спикеров
            self.log("2. Анализ спикеров...")
            self.update_progress(30)
            speaker_segments = self.perform_speaker_diarization(temp_audio_path)
            
            # Шаг 3: Транскрибация
            self.log("3. Транскрибация аудио...")
            self.update_progress(50)
            transcription = self.transcribe_audio(temp_audio_path, language)
            
            if transcription is None:
                return False
            
            # Шаг 4: Объединение результатов
            self.log("4. Объединение результатов...")
            self.update_progress(80)
            merged_result = self.merge_transcription_with_speakers(
                transcription, speaker_segments
            )
            
            # Шаг 5: Создание документа
            self.log("5. Создание документа...")
            self.update_progress(90)
            output_path = output_dir / f"{video_path.stem}_транскрибация.docx"
            self.create_word_document(merged_result, str(output_path), video_path.name)
            
            self.update_progress(100)
            self.log("Обработка завершена успешно!")
            return True
            
        except Exception as e:
            self.log(f"Ошибка обработки: {e}")
            return False
        finally:
            if os.path.exists(temp_audio_path):
                os.unlink(temp_audio_path)


class TranscriberGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Транскрибер видео")
        self.root.geometry("800x600")
        
        # Переменные
        self.video_path = tk.StringVar()
        self.output_path = tk.StringVar()
        self.language = tk.StringVar(value="ru")
        self.whisper_model = tk.StringVar(value="base")
        self.processing = False
        
        self.setup_ui()
        
    def setup_ui(self):
        """Создание интерфейса"""
        # Главный фрейм
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Настройка растягивания
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        
        # Заголовок
        title_label = ttk.Label(main_frame, text="🎥 Транскрибер видео", font=("Arial", 16, "bold"))
        title_label.grid(row=0, column=0, columnspan=3, pady=(0, 20))
        
        # Выбор видео файла
        ttk.Label(main_frame, text="Видео файл:").grid(row=1, column=0, sticky=tk.W, pady=5)
        ttk.Entry(main_frame, textvariable=self.video_path, width=50).grid(row=1, column=1, sticky=(tk.W, tk.E), padx=(5, 5))
        ttk.Button(main_frame, text="Обзор...", command=self.browse_video).grid(row=1, column=2, padx=(5, 0))
        
        # Выбор папки для сохранения
        ttk.Label(main_frame, text="Папка сохранения:").grid(row=2, column=0, sticky=tk.W, pady=5)
        ttk.Entry(main_frame, textvariable=self.output_path, width=50).grid(row=2, column=1, sticky=(tk.W, tk.E), padx=(5, 5))
        ttk.Button(main_frame, text="Обзор...", command=self.browse_output).grid(row=2, column=2, padx=(5, 0))
        
        # Настройки
        settings_frame = ttk.LabelFrame(main_frame, text="Настройки", padding="10")
        settings_frame.grid(row=3, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=10)
        settings_frame.columnconfigure(1, weight=1)
        
        # Язык
        ttk.Label(settings_frame, text="Язык:").grid(row=0, column=0, sticky=tk.W)
        language_combo = ttk.Combobox(settings_frame, textvariable=self.language, width=15)
        language_combo['values'] = ('ru', 'en', 'auto')
        language_combo.grid(row=0, column=1, sticky=tk.W, padx=(5, 20))
        
        # Модель Whisper
        ttk.Label(settings_frame, text="Модель Whisper:").grid(row=0, column=2, sticky=tk.W)
        model_combo = ttk.Combobox(settings_frame, textvariable=self.whisper_model, width=15)
        model_combo['values'] = ('tiny', 'base', 'small', 'medium', 'large')
        model_combo.grid(row=0, column=3, sticky=tk.W, padx=(5, 0))
        
        # Прогресс бар
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(main_frame, variable=self.progress_var, maximum=100)
        self.progress_bar.grid(row=4, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=10)
        
        # Статус
        self.status_var = tk.StringVar(value="Готов к работе")
        status_label = ttk.Label(main_frame, textvariable=self.status_var)
        status_label.grid(row=5, column=0, columnspan=3, pady=5)
        
        # Кнопки
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=6, column=0, columnspan=3, pady=10)
        
        self.start_button = ttk.Button(button_frame, text="🚀 Начать обработку", command=self.start_processing)
        self.start_button.pack(side=tk.LEFT, padx=5)
        
        ttk.Button(button_frame, text="📁 Открыть папку результатов", command=self.open_output_folder).pack(side=tk.LEFT, padx=5)
        
        ttk.Button(button_frame, text="ℹ️ О программе", command=self.show_about).pack(side=tk.LEFT, padx=5)
        
        # Лог
        log_frame = ttk.LabelFrame(main_frame, text="Журнал выполнения", padding="5")
        log_frame.grid(row=7, column=0, columnspan=3, sticky=(tk.W, tk.E, tk.N, tk.S), pady=10)
        log_frame.columnconfigure(0, weight=1)
        log_frame.rowconfigure(0, weight=1)
        main_frame.rowconfigure(7, weight=1)
        
        self.log_text = scrolledtext.ScrolledText(log_frame, width=80, height=15)
        self.log_text.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
    def browse_video(self):
        """Выбор видео файла"""
        filename = filedialog.askopenfilename(
            title="Выберите видео файл",
            filetypes=[
                ("Видео файлы", "*.mp4 *.avi *.mov *.mkv *.webm *.flv *.wmv"),
                ("Все файлы", "*.*")
            ]
        )
        if filename:
            self.video_path.set(filename)
            # Автоматически устанавливаем папку сохранения
            if not self.output_path.get():
                self.output_path.set(os.path.dirname(filename))
    
    def browse_output(self):
        """Выбор папки для сохранения"""
        dirname = filedialog.askdirectory(title="Выберите папку для сохранения")
        if dirname:
            self.output_path.set(dirname)
    
    def log_message(self, message):
        """Добавление сообщения в лог"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.log_text.insert(tk.END, f"[{timestamp}] {message}\n")
        self.log_text.see(tk.END)
        self.root.update()
    
    def update_progress(self, value):
        """Обновление прогресс-бара"""
        self.progress_var.set(value)
        self.root.update()
    
    def update_status(self, status):
        """Обновление статуса"""
        self.status_var.set(status)
        self.root.update()
    
    def start_processing(self):
        """Запуск обработки в отдельном потоке"""
        if self.processing:
            return
            
        if not self.video_path.get():
            messagebox.showerror("Ошибка", "Выберите видео файл")
            return
        
        if not os.path.exists(self.video_path.get()):
            messagebox.showerror("Ошибка", "Видео файл не найден")
            return
        
        self.processing = True
        self.start_button.config(text="⏳ Обработка...", state="disabled")
        self.progress_var.set(0)
        self.log_text.delete(1.0, tk.END)
        
        # Запуск в отдельном потоке
        thread = threading.Thread(target=self.process_video_thread)
        thread.daemon = True
        thread.start()
    
    def process_video_thread(self):
        """Обработка видео в отдельном потоке"""
        try:
            self.update_status("Инициализация...")
            
            transcriber = VideoTranscriber(
                progress_callback=self.update_progress,
                log_callback=self.log_message
            )
            
            # Изменяем модель если нужно
            if self.whisper_model.get() != "base":
                self.log_message(f"Загрузка модели {self.whisper_model.get()}...")
                transcriber.whisper_model = whisper.load_model(self.whisper_model.get())
            
            self.update_status("Обработка видео...")
            
            success = transcriber.process_video(
                self.video_path.get(),
                self.output_path.get() if self.output_path.get() else None,
                self.language.get()
            )
            
            if success:
                self.update_status("✅ Обработка завершена успешно!")
                messagebox.showinfo("Успех", "Транскрибация завершена!\nДокument сохранен в выбранной папке.")
            else:
                self.update_status("❌ Ошибка при обработке")
                messagebox.showerror("Ошибка", "Произошла ошибка при обработке видео.\nПроверьте журнал для подробностей.")
                
        except Exception as e:
            self.log_message(f"Критическая ошибка: {e}")
            self.update_status("❌ Критическая ошибка")
            messagebox.showerror("Критическая ошибка", f"Произошла критическая ошибка:\n{e}")
        
        finally:
            self.processing = False
            self.start_button.config(text="🚀 Начать обработку", state="normal")
    
    def open_output_folder(self):
        """Открытие папки с результатами"""
        output_dir = self.output_path.get() or os.path.dirname(self.video_path.get())
        if output_dir and os.path.exists(output_dir):
            if sys.platform == "win32":
                os.startfile(output_dir)
            elif sys.platform == "darwin":
                subprocess.Popen(["open", output_dir])
            else:
                subprocess.Popen(["xdg-open", output_dir])
        else:
            messagebox.showwarning("Предупреждение", "Папка не найдена")
    
    def show_about(self):
        """Показать информацию о программе"""
        about_text = """
Транскрибер видео v1.0

Программа для автоматической транскрибации видео 
с разделением по спикерам.

Возможности:
• Поддержка множества видео форматов
• Разделение спикеров (до 10 человек)
• Русский и английский языки
• Экспорт в Word документ
• Временные метки

Технологии:
• OpenAI Whisper - транскрибация
• Pyannote Audio - разделение спикеров
• FFmpeg - обработка видео

Разработано с ❤️
        """
        messagebox.showinfo("О программе", about_text)


def main():
    root = tk.Tk()
    app = TranscriberGUI(root)
    root.mainloop()


if __name__ == "__main__":
    main()